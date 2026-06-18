import os
import random
import requests
import numpy as np
import pandas as pd
from sklearn.preprocessing import StandardScaler
from sklearn.ensemble import IsolationForest
import torch
import torch.nn as nn
from sentence_transformers import SentenceTransformer

SEED = 42
np.random.seed(SEED)
random.seed(SEED)
torch.manual_seed(SEED)

os.makedirs("data", exist_ok=True)

ERP_NEXT_BRANCH = "develop"
ERP_NEXT_OWNER = "frappe"
ERP_NEXT_REPO = "erpnext"


def collect_erpnext_artifacts(branch=ERP_NEXT_BRANCH, max_files=40):
    print("[Agent 1A] Connecting to ERPNext GitHub repository...")
    accounting_keywords = [
        "account", "accounts", "accounting", "journal", "journal_entry",
        "general_ledger", "gl_entry", "ledger", "invoice", "sales_invoice",
        "purchase_invoice", "payment_entry", "payment", "bank", "banking",
        "tax", "company", "cost_center", "asset", "liability", "revenue", "expense"
    ]
    tree_url = f"https://api.github.com/repos/{ERP_NEXT_OWNER}/{ERP_NEXT_REPO}/git/trees/{branch}?recursive=1"
    headers = {}
    token = os.environ.get("GITHUB_TOKEN")
    if token:
        headers["Authorization"] = f"token {token}"
    try:
        response = requests.get(tree_url, headers=headers, timeout=25)
        response.raise_for_status()
        tree_items = response.json().get("tree", [])
        candidate_paths = []
        for item in tree_items:
            path = item.get("path", "")
            item_type = item.get("type", "")
            lower_path = path.lower()
            if item_type != "blob":
                continue
            if not lower_path.endswith((".json", ".csv", ".py", ".md", ".txt")):
                continue
            if any(keyword in lower_path for keyword in accounting_keywords):
                candidate_paths.append(path)
        candidate_paths = candidate_paths[:max_files]
        records = []
        for path in candidate_paths:
            raw_url = f"https://raw.githubusercontent.com/{ERP_NEXT_OWNER}/{ERP_NEXT_REPO}/{branch}/{path}"
            detected = sorted({kw for kw in accounting_keywords if kw in path.lower()})
            artifact_type = "ERPNext source artifact"
            if path.lower().endswith(".json"):
                artifact_type = "ERPNext JSON DocType/configuration artifact"
            elif path.lower().endswith(".csv"):
                artifact_type = "ERPNext CSV fixture/tabular artifact"
            elif path.lower().endswith(".py"):
                artifact_type = "ERPNext Python business logic/test artifact"
            elif path.lower().endswith(".md"):
                artifact_type = "ERPNext documentation artifact"
            records.append({
                "repository": f"{ERP_NEXT_OWNER}/{ERP_NEXT_REPO}",
                "branch": branch,
                "source_path": path,
                "raw_url": raw_url,
                "artifact_type": artifact_type,
                "detected_keywords": ", ".join(detected)
            })
        artifacts_df = pd.DataFrame(records)
        if artifacts_df.empty:
            print("[Agent 1A] No ERPNext artifacts found. Fallback metadata will be used.")
        else:
            artifacts_df.to_csv("data/erpnext_github_artifacts.csv", index=False)
            print(f"[Agent 1A] ERPNext artifacts collected: {len(artifacts_df)}")
            print("[Agent 1A] Saved: data/erpnext_github_artifacts.csv")
        return artifacts_df
    except Exception as e:
        print(f"[Agent 1A] ERPNext GitHub ingestion failed: {e}")
        return pd.DataFrame()


def generate_research_grade_ledger(artifacts_df=None, samples=15):
    np.random.seed(SEED)
    companies = [f"Global_Enterprise_{100 + i}" for i in range(samples)]
    debit = np.random.uniform(50000, 2500000, samples)
    credit = debit.copy()
    credit[3] = credit[3] * 1.35
    credit[8] = credit[8] * 0.60
    procurement_waste_tons = np.random.uniform(5, 120, samples)
    procurement_waste_tons[12] = 450.0
    financial_narratives = [
        "Revenue recognized through normal channels matching invoice schedules flawlessly.",
        "Period-end financial close operations executed in accordance with IFRS frameworks.",
        "Standard amortization of intangible patent portfolios calculated across linear schedules.",
        "Special batch ledger adjustments committed via root administrative bypass authorization.",
        "Fixed asset reconciliation completed successfully against factory logistics tracking indices.",
        "Accounts receivable collections traced back securely to third-party merchant processor accounts.",
        "Routine accrual adjustments for utility overhead metrics performed under corporate supervision.",
        "Tax provision balancing entry executed following verified regional updates.",
        "Suspense account balancing clearing entry committed without secondary supervisory sign-off.",
        "Foreign exchange translation calculation matching real-time valuation windows.",
        "Deferred revenue segments released systematically following completion of client milestones.",
        "Intercompany balance netting process finalized across standard regional ledgers.",
        "Provisions for inventory obsolescence written down according to warehouse audits.",
        "Prepaid operating lease allocations systematically distributed to local business accounts.",
        "Minor adjustment to payroll expense variances matching labor output reports."
    ]
    esg_disclosures = [
        "Scope 1 emissions lowered by 14% through continuous solar infrastructure integration.",
        "Corporate inclusion goals hit target thresholds across operational manufacturing plants.",
        "Waste management protocols met all municipal environmental safety compliance rules.",
        "Supply chain validation audit verified zero exposure to child labor vectors.",
        "Corporate literature states our factory operation has achieved a net-zero impact framework instantly.",
        "Board restructuring enhanced transparency via addition of two independent monitoring leads.",
        "Strategic sourcing pipelines verified that 80% of primary components use recycled compounds.",
        "Employee training hours on compliance topics expanded by 25% year-over-year.",
        "Water remediation investments completed across processing facility footprints.",
        "Donation and lobbying records published completely online to ensure corporate accountability.",
        "Workplace protection metrics registered zero high-severity incidents for this cycle.",
        "Sustainable sourcing guidelines formalized across all tier-one logistics providers.",
        "Public report affirms absolute ecological care while heavy localized waste spikes continue.",
        "Community engagement initiatives deployed across new infrastructure building sites.",
        "Data privacy protection standards updated to enforce global regulatory compliance."
    ]
    cyber_disclosures = [
        "Perimeter logging metrics confirm zero security bypass actions or identity theft instances.",
        "Enterprise access points migrated securely to multi-factor zero-trust validation models.",
        "Phishing testing exercises achieved exceptional compliance values among ledger accountants.",
        "Database schemas for customer records encrypted via standard cryptographic layers.",
        "Critical ERP patch updates delayed due to legacy hardware optimization blocks.",
        "Regular perimeter scanning validated effective isolation of payment gateway components.",
        "Identity permission vectors verified strictly via role-based access control paradigms.",
        "Disaster recovery backups generated nightly and mirrored to cloud locations.",
        "Incident containment runbooks tested and approved by independent external panels.",
        "Network load monitors confirmed normal performance metrics throughout this cycle.",
        "Endpoint configuration auditing tools confirmed clean operation for core accounting infrastructure.",
        "Third-party system risk management protocols updated to include vendor tracking arrays.",
        "Minor firewall policy adjustments completed to isolate development testbeds.",
        "Access key rotation matrices executed across production database instances.",
        "Data governance rules updated to meet evolving multi-regional security rules."
    ]
    if artifacts_df is not None and not artifacts_df.empty:
        paths = artifacts_df["source_path"].tolist()
        types = artifacts_df["artifact_type"].tolist()
        keywords = artifacts_df["detected_keywords"].tolist()
    else:
        paths = ["ERPNext fallback simulator"]
        types = ["Synthetic ERP ledger simulator"]
        keywords = ["ledger, invoice, accounting"]
    df = pd.DataFrame({
        "Company_ID": companies,
        "Account_Segment": np.random.choice(["1000 Assets", "4000 Revenue", "2000 Liabilities", "5000 Expenses"], samples),
        "Debit_Value": debit,
        "Credit_Value": credit,
        "Procurement_Waste_Tons": procurement_waste_tons,
        "Financial_Narrative": financial_narratives,
        "ESG_Disclosure": esg_disclosures,
        "Cyber_Disclosure": cyber_disclosures,
        "ERPNext_Source_Path": [paths[i % len(paths)] for i in range(samples)],
        "ERPNext_Artifact_Type": [types[i % len(types)] for i in range(samples)],
        "ERPNext_Detected_Keywords": [keywords[i % len(keywords)] for i in range(samples)]
    })
    df.to_csv("data/raw_erp_accounting_data.csv", index=False)
    print("[Agent 1B] Synthetic ERP/GL matrix saved: data/raw_erp_accounting_data.csv")
    return df


def preprocessing_agent(df):
    cleaned_df = df.drop_duplicates().ffill().bfill().copy()
    cleaned_df["Balance_Variance"] = (cleaned_df["Debit_Value"] - cleaned_df["Credit_Value"]).abs()
    numeric_cols = ["Debit_Value", "Credit_Value", "Procurement_Waste_Tons", "Balance_Variance"]
    scaler = StandardScaler()
    scaled_array = scaler.fit_transform(cleaned_df[numeric_cols])
    for i, col in enumerate(numeric_cols):
        cleaned_df[f"scaled_{col}"] = scaled_array[:, i]
    print("[Agent 2] Preprocessing complete.")
    return cleaned_df


def keyword_score(text, terms):
    text = str(text).lower()
    matches = sum(1 for term in terms if term in text)
    return min(matches / max(len(terms), 1), 1.0)


def multi_bert_agent(df):
    df = df.copy()
    print("[Agent 3] Loading SentenceTransformer model...")
    embedder = SentenceTransformer("all-MiniLM-L6-v2")
    combined_text = (
        df["Financial_Narrative"].astype(str) + " " +
        df["ESG_Disclosure"].astype(str) + " " +
        df["Cyber_Disclosure"].astype(str) + " " +
        df["ERPNext_Source_Path"].astype(str)
    ).tolist()
    embeddings = embedder.encode(combined_text, show_progress_bar=False)
    ledger_terms = ["bypass", "suspense", "manual", "adjustment", "root", "without secondary", "period-end"]
    cyber_terms = ["patch", "legacy", "bypass", "identity", "access", "security", "firewall", "encryption"]
    esg_terms = ["waste", "net-zero", "emissions", "ecological", "sustainable", "recycled", "green", "privacy"]
    df["SubAgent_A_GeneralBERT_Risk"] = np.clip(np.var(embeddings, axis=1), 0, 1)
    df["SubAgent_B_LedgerBERT_Risk"] = df["Financial_Narrative"].apply(lambda x: keyword_score(x, ledger_terms))
    df["SubAgent_C_CyberBERT_Risk"] = df["Cyber_Disclosure"].apply(lambda x: keyword_score(x, cyber_terms))
    df["SubAgent_D_ESGBERT_Risk"] = df["ESG_Disclosure"].apply(lambda x: keyword_score(x, esg_terms))
    df["Text_Embedding_Proxy"] = np.mean(embeddings, axis=1)
    print("[Agent 3] Multi-BERT proxy scoring complete.")
    return df, embeddings


class PyTorchAutoencoder(nn.Module):
    def __init__(self, input_dim):
        super().__init__()
        self.encoder = nn.Sequential(
            nn.Linear(input_dim, 6),
            nn.Tanh(),
            nn.Linear(6, 3)
        )
        self.decoder = nn.Sequential(
            nn.Linear(3, 6),
            nn.Tanh(),
            nn.Linear(6, input_dim)
        )

    def forward(self, x):
        return self.decoder(self.encoder(x))


def autoencoder_anomaly_agent(df, threshold_percentile=85, epochs=160, lr=0.03):
    df = df.copy()
    torch.manual_seed(SEED)
    feature_cols = [
        "scaled_Debit_Value",
        "scaled_Credit_Value",
        "scaled_Procurement_Waste_Tons",
        "scaled_Balance_Variance",
        "SubAgent_A_GeneralBERT_Risk",
        "SubAgent_B_LedgerBERT_Risk",
        "SubAgent_C_CyberBERT_Risk",
        "SubAgent_D_ESGBERT_Risk",
        "Text_Embedding_Proxy"
    ]
    X = df[feature_cols].astype(float).values
    X_tensor = torch.tensor(X, dtype=torch.float32)
    model = PyTorchAutoencoder(input_dim=len(feature_cols))
    criterion = nn.MSELoss()
    optimizer = torch.optim.Adam(model.parameters(), lr=lr)
    model.train()
    for epoch in range(epochs):
        optimizer.zero_grad()
        outputs = model(X_tensor)
        loss = criterion(outputs, X_tensor)
        loss.backward()
        optimizer.step()
    model.eval()
    with torch.no_grad():
        reconstructed = model(X_tensor)
        errors = torch.mean((X_tensor - reconstructed) ** 2, dim=1).numpy()
    threshold_value = np.percentile(errors, threshold_percentile)
    df["Reconstruction_Error"] = errors
    df["Autoencoder_Threshold"] = threshold_value
    df["Anomaly_Indicator"] = df["Reconstruction_Error"].apply(
        lambda e: "CRITICAL PROFILE ALERT" if e > threshold_value else "Normal Verified Ledger"
    )
    print("[Agent 4A] PyTorch Autoencoder anomaly detection complete.")
    return df


def isolation_forest_agent(df, text_embeddings, contamination=0.30):
    df = df.copy()
    structured_cols = [
        "scaled_Debit_Value",
        "scaled_Credit_Value",
        "scaled_Procurement_Waste_Tons",
        "scaled_Balance_Variance",
        "SubAgent_B_LedgerBERT_Risk",
        "SubAgent_C_CyberBERT_Risk",
        "SubAgent_D_ESGBERT_Risk"
    ]
    structured_features = df[structured_cols].astype(float).values
    if text_embeddings is not None and len(text_embeddings) == len(df):
        combined_features = np.hstack([text_embeddings, structured_features])
    else:
        print("Warning: text_embeddings not available or length mismatch. Using only structured features for Isolation Forest.")
        combined_features = structured_features
    model = IsolationForest(contamination=contamination, random_state=SEED)
    pred = model.fit_predict(combined_features)
    df["IForest_Anomaly"] = [1 if p == -1 else 0 for p in pred]
    df["IForest_Indicator"] = df["IForest_Anomaly"].apply(
        lambda x: "CRITICAL PROFILE ALERT" if x == 1 else "Normal Verified Ledger"
    )
    print("[Agent 4B] Isolation Forest fallback complete.")
    return df


def esg_risk_agent(df):
    df = df.copy()
    waste = df["scaled_Procurement_Waste_Tons"].astype(float)
    waste_norm = (waste - waste.min()) / (waste.max() - waste.min() + 1e-6)
    df["Environmental_Score"] = (100 * (1 - waste_norm)).clip(15, 98)
    df["Social_Score"] = (100 * (1 - df["SubAgent_C_CyberBERT_Risk"].astype(float))).clip(20, 99)
    df["Governance_Score"] = (100 * (1 - df["SubAgent_B_LedgerBERT_Risk"].astype(float))).clip(10, 97)
    df["Overall_ESG_Score"] = (
        df["Environmental_Score"] + df["Social_Score"] + df["Governance_Score"]
    ) / 3
    def classify(row):
        if row["Anomaly_Indicator"] == "CRITICAL PROFILE ALERT" or row["Overall_ESG_Score"] < 50:
            return "High"
        elif row["Overall_ESG_Score"] < 75 or row["IForest_Indicator"] == "CRITICAL PROFILE ALERT":
            return "Medium"
        return "Low"
    df["ESG_Risk_Classification"] = df.apply(classify, axis=1)
    print("[Agent 5] ESG risk assessment complete.")
    return df


def compliance_agent(df):
    df = df.copy()
    def score(row):
        compliance = 100
        if row["Anomaly_Indicator"] == "CRITICAL PROFILE ALERT":
            compliance -= 30
        if row["IForest_Indicator"] == "CRITICAL PROFILE ALERT":
            compliance -= 15
        if row["SubAgent_B_LedgerBERT_Risk"] > 0.20:
            compliance -= 20
        if row["SubAgent_C_CyberBERT_Risk"] > 0.20:
            compliance -= 15
        if row["Overall_ESG_Score"] < 60:
            compliance -= 10
        return max(compliance, 0)
    df["EU_AI_Act_Compliance_Score"] = df.apply(score, axis=1)
    def log(row):
        reasons = []
        if row["Anomaly_Indicator"] == "CRITICAL PROFILE ALERT":
            reasons.append(f"high autoencoder reconstruction loss ({row['Reconstruction_Error']:.4f})")
        if row["IForest_Indicator"] == "CRITICAL PROFILE ALERT":
            reasons.append("Isolation Forest fallback also flagged this profile")
        if row["SubAgent_B_LedgerBERT_Risk"] > 0.20:
            reasons.append("ledger-risk terminology was detected")
        if row["SubAgent_C_CyberBERT_Risk"] > 0.20:
            reasons.append("cybersecurity-risk terminology was detected")
        if row["Overall_ESG_Score"] < 60:
            reasons.append("ESG score is below the review threshold")
        if not reasons:
            return "Transparent automated classification. No immediate human escalation required under prototype thresholds."
        return "Human verification recommended because " + "; ".join(reasons) + "."
    df["AI_Explainability_Log"] = df.apply(log, axis=1)
    print("[Agent 6] Trustworthy AI compliance audit complete.")
    return df


import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches


def dashboard_agent(df):
    print("Dashboard is disabled in this runner.")


def export_all_tables_to_excel(
    artifacts,
    raw,
    preprocessed,
    bert_enriched,
    ae_scored,
    if_scored,
    esg_scored,
    final,
    excel_path="data/multi_agent_accounting_ai_results.xlsx"
):
    with pd.ExcelWriter(excel_path, engine="openpyxl") as writer:
        artifacts.to_excel(writer, sheet_name="erpnext_artifacts", index=False)
        raw.to_excel(writer, sheet_name="raw_erp_ledger", index=False)
        preprocessed.to_excel(writer, sheet_name="preprocessed", index=False)
        bert_enriched.to_excel(writer, sheet_name="bert_enriched", index=False)
        ae_scored.to_excel(writer, sheet_name="autoencoder_scored", index=False)
        if_scored.to_excel(writer, sheet_name="isolation_forest_scored", index=False)
        esg_scored.to_excel(writer, sheet_name="esg_scored", index=False)
        final.to_excel(writer, sheet_name="final_output", index=False)
    print(f"Excel workbook saved to: {excel_path}")
    return excel_path


def export_figures_to_word(df, word_path="data/multi_agent_accounting_ai_figures.docx"):
    doc = Document()
    doc.add_heading('Multi-Agent Accounting AI Figures', level=1)

    # Figure 1: Reconstruction Error by Company
    fig, ax = plt.subplots(figsize=(10, 5))
    ax.bar(df['Company_ID'], df['Reconstruction_Error'], color='steelblue')
    ax.set_title('Reconstruction Error by Company')
    ax.set_xlabel('Company_ID')
    ax.set_ylabel('Reconstruction Error')
    ax.set_xticklabels(df['Company_ID'], rotation=45, ha='right')
    fig.tight_layout()
    fig_path1 = 'data/figure_reconstruction_error.png'
    fig.savefig(fig_path1)
    plt.close(fig)
    doc.add_heading('Reconstruction Error by Company', level=2)
    doc.add_picture(fig_path1, width=Inches(6))

    # Figure 2: Overall ESG Score by Company
    fig, ax = plt.subplots(figsize=(10, 5))
    ax.bar(df['Company_ID'], df['Overall_ESG_Score'], color='seagreen')
    ax.set_title('Overall ESG Score by Company')
    ax.set_xlabel('Company_ID')
    ax.set_ylabel('Overall ESG Score')
    ax.set_xticklabels(df['Company_ID'], rotation=45, ha='right')
    fig.tight_layout()
    fig_path2 = 'data/figure_esg_score.png'
    fig.savefig(fig_path2)
    plt.close(fig)
    doc.add_heading('Overall ESG Score by Company', level=2)
    doc.add_picture(fig_path2, width=Inches(6))

    # Figure 3: EU AI Act Compliance Score by Company
    fig, ax = plt.subplots(figsize=(10, 5))
    ax.bar(df['Company_ID'], df['EU_AI_Act_Compliance_Score'], color='indianred')
    ax.set_title('EU AI Act Compliance Score by Company')
    ax.set_xlabel('Company_ID')
    ax.set_ylabel('Compliance Score')
    ax.set_xticklabels(df['Company_ID'], rotation=45, ha='right')
    fig.tight_layout()
    fig_path3 = 'data/figure_compliance_score.png'
    fig.savefig(fig_path3)
    plt.close(fig)
    doc.add_heading('EU AI Act Compliance Score by Company', level=2)
    doc.add_picture(fig_path3, width=Inches(6))

    doc.save(word_path)
    print(f"Word document saved to: {word_path}")
    return word_path


def run_multi_agent_accounting_ai(branch=ERP_NEXT_BRANCH, max_files=40, show_dashboard=False, output_path="data/final_agent_output.csv", excel_output_path="data/multi_agent_accounting_ai_results.xlsx", word_output_path="data/multi_agent_accounting_ai_figures.docx"):
    print("\n" + "=" * 76)
    print("STARTING MULTI-AGENT ACCOUNTING AI — RUNNER")
    print("=" * 76 + "\n")
    artifacts = collect_erpnext_artifacts(branch=branch, max_files=max_files)
    raw = generate_research_grade_ledger(artifacts)
    preprocessed = preprocessing_agent(raw)
    bert_enriched, embeddings = multi_bert_agent(preprocessed)
    ae_scored = autoencoder_anomaly_agent(bert_enriched)
    if_scored = isolation_forest_agent(ae_scored, embeddings)
    esg_scored = esg_risk_agent(if_scored)
    final = compliance_agent(esg_scored)
    final = final.drop(columns=["Text_Embedding_Proxy"], errors="ignore")
    final.to_csv(output_path, index=False)
    excel_path = export_all_tables_to_excel(
        artifacts,
        raw,
        preprocessed,
        bert_enriched,
        ae_scored,
        if_scored,
        esg_scored,
        final,
        excel_path=excel_output_path,
    )
    word_path = export_figures_to_word(final, word_path=word_output_path)
    print(f"\nFinal output saved to: {output_path}")
    print(f"Excel workbook saved to: {excel_path}")
    print(f"Word document saved to: {word_path}")
    if show_dashboard:
        dashboard_agent(final)
    print("\n" + "=" * 76)
    print("WORKFLOW COMPLETE")
    print("=" * 76 + "\n")
    return final


if __name__ == "__main__":
    result = run_multi_agent_accounting_ai(show_dashboard=False)
    print("Result rows:", len(result) if hasattr(result, "__len__") else "?")
